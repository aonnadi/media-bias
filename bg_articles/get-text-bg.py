import io
import re
import csv
import os.path
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from docx import Document

'''
This code is to extract text from Google Drive links in a CSV file.
'''
# If modifying these scopes, delete the file token.json.
SCOPES = ["https://www.googleapis.com/auth/drive.readonly"]

def authenticate_google_drive_api():
    """Authenticate and return a Google Drive service object."""
    creds = None
    if os.path.exists("token.json"):
        creds = Credentials.from_authorized_user_file("token.json", SCOPES)
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('../credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        with open("token.json", 'w') as token:
            token.write(creds.to_json())

    return build('drive', 'v3', credentials=creds)

def download_and_extract_text_from_docx(drive_service, file_id):
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while not done:
        _, done = downloader.next_chunk()

    fh.seek(0)
    document = Document(fh)
    
    full_text = []

    # Extract from the document body
    for para in document.paragraphs:
        full_text.append(para.text)
    
    # Attempt to extract from headers
    for section in document.sections:
        for header in section.header.paragraphs:
            full_text.append(header.text)
    
    # Attempt to extract from footers
    for section in document.sections:
        for footer in section.footer.paragraphs:
            full_text.append(footer.text)

    # Combine all text into a single string
    return '\n'.join(full_text)

def extract_text_from_link(drive_service, link):
    """Extract text from given Google Drive link."""
    # Regex for extracting file ID from Google Drive link
    file_id_match = re.search(r'(?:/d/|id=)([^/&?]+)', link)
    if not file_id_match:
        print(f"Could not extract file ID from link: {link}")
        return None

    file_id = file_id_match.group(1)

    try:
        # Get the file's MIME type
        file_metadata = drive_service.files().get(fileId=file_id, fields='mimeType').execute()
        mime_type = file_metadata.get('mimeType')

        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            # Download and extract text from .docx file
            return download_and_extract_text_from_docx(drive_service, file_id)
        else:
            print(f"Unsupported MIME type: {mime_type}")
            return None
    except Exception as e:
        print(f"An error occurred: {e}")
        return None

def main():
    drive_service = authenticate_google_drive_api()

    # Replace with the path to CSV file containing Google Drive links
    csv_file_path = '../datasets/labeling-responses.csv'
    output_csv_path = 'tbg_results.csv'  # Define the output CSV file path

    # Initialize a list to hold the results
    results = []

    # Extract Google Drive links from CSV
    with open(csv_file_path, mode='r', encoding='utf-8') as file:
        reader = csv.reader(file)
        for row in reader:
            for cell in row:
                links = re.findall(r'https://drive\.google.com/[^ ]+', cell)
                for link in links:
                    text = extract_text_from_link(drive_service, link)
                    if text:
                        # For demonstration, only save the first 100 characters of the extracted text
                        results.append([text, link, 'Race Discussed Placeholder'])
                        print(f"Extracted text from {link}")
                    else:
                        results.append(['Failed to extract text or unsupported file type', link, 'Race Discussed Placeholder'])
                        print(f"Failed to extract text from {link}.")

    # Write the results to a new CSV file
    with open(output_csv_path, mode='w', newline='', encoding='utf-8') as output_file:
        writer = csv.writer(output_file)
        writer.writerow(['text', 'link', 'race_label', 'race_discussed'])  # Writing the headers
        writer.writerows(results)  # Writing the extracted data

    print(f"Processed {len(results)} files. Results saved to {output_csv_path}.")

if __name__ == "__main__":
    main()

